#!/usr/bin/env python3
"""
Build a universal GyanQuest competition project report (DOCX).

Examples:
  py -3 tools/build_project_report.py
  py -3 tools/build_project_report.py --team "Team Alpha" --members "Md Tahsan Islam, Mufrid Johanee, Shadiya Zaman Tanha, Md Mohaimenul Islam"
  py -3 tools/build_project_report.py --event "National STEM Fair 2026" --institution "Your college"
"""
from __future__ import annotations

import argparse
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "GyanQuest_Project_Report.docx"

INK = RGBColor(0x1B, 0x28, 0x38)
MUTED = RGBColor(0x5C, 0x6B, 0x7A)
HEADING = RGBColor(0x1E, 0x4D, 0x6B)
GOLD = RGBColor(0x9A, 0x7B, 0x3C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
IVORY = RGBColor(0xF4, 0xEB, 0xD0)
COVER_SUB = RGBColor(0xD7, 0xE0, 0xEA)
FILL_HINT = "F7F1E3"
NAVY = "1E3A5F"
GOLD_HEX = "C4A35A"
ROW_ALT = "F4F7FA"
VALUE_FILL = "EEF3F8"


def shade(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_run(run, *, size=11, bold=False, color=INK, italic=False, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_p(doc, text, *, size=11, bold=False, color=INK, align="left", space_after=8, space_before=0, italic=False):
    p = doc.add_paragraph()
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color, italic=italic)
    return p


def add_h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = HEADING if level == 1 else INK
        run.font.name = "Calibri"
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        for run in p.runs:
            set_run(run, size=11)


def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        p.paragraph_format.space_after = Pt(3)
        for run in p.runs:
            set_run(run, size=11)


def add_table(doc, headers, rows, col_widths=None, header_fill=NAVY):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_run(run, size=10, bold=True, color=WHITE)
        shade(cell, header_fill)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run(run, size=10)
            if r_i % 2 == 1:
                shade(cell, ROW_ALT)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def parse_members(raw: str) -> list[tuple[str, str]]:
    out = []
    for part in raw.split(","):
        part = part.strip()
        if not part:
            continue
        if ":" in part:
            name, role = part.split(":", 1)
            out.append((name.strip(), role.strip()))
        else:
            out.append((part, "Team member"))
    return out


def build(args) -> Path:
    team = (args.team or "").strip() or "[Team name]"
    event = (args.event or "").strip() or "[Competition / event name]"
    institution = (args.institution or "").strip() or "[Institution / organization]"
    members = parse_members(args.members) if args.members.strip() else [
        ("[Member 1 full name]", "[Role]"),
        ("[Member 2 full name]", "[Role]"),
        ("[Member 3 full name]", "[Role]"),
        ("[Member 4 full name]", "[Role]"),
    ]
    placeholder = team.startswith("[") or event.startswith("[") or members[0][0].startswith("[")
    today = date.today().strftime("%d %B %Y")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("GyanQuest  ·  ImpactX  ·  Project Report")
    set_run(hr, size=9, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Confidential to the submitting team  ·  Page ")
    set_run(fr, size=9, color=MUTED)
    # PAGE field
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    fp.add_run()._element.append(fld)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fp.add_run()._element.append(instr)
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    fp.add_run()._element.append(fld2)

    # ----- Cover -----
    banner = doc.add_table(rows=2, cols=1)
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = banner.rows[0].cells[0]
    shade(cell, NAVY)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("GYANQUEST")
    set_run(r, size=28, bold=True, color=IVORY, name="Calibri")
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run("Project Report")
    set_run(r, size=16, bold=True, color=WHITE)
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p3.add_run("Interactive bilingual mission games for school STEM  ·  Canvas 2D curriculum  ·  3D Specimen Lab")
    set_run(r, size=11, color=COVER_SUB)
    p4 = cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_after = Pt(14)
    r = p4.add_run("Offline-first  ·  English + Bangla  ·  Browser, not a hardware lab")
    set_run(r, size=11, color=IVORY)
    gold_bar = banner.rows[1].cells[0]
    shade(gold_bar, GOLD_HEX)
    gp = gold_bar.paragraphs[0]
    gp.paragraph_format.space_before = Pt(2)
    gp.paragraph_format.space_after = Pt(2)
    gr = gp.add_run(" ")
    set_run(gr, size=4, color=WHITE)

    add_p(doc, "", space_after=6)

    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    fields = [
        ("Competition / event", event),
        ("Team name", team),
        ("Institution", institution),
        ("Submission date", today),
    ]
    for i, (lab, val) in enumerate(fields):
        c0, c1 = meta.rows[i].cells
        c0.text = ""
        c1.text = ""
        shade(c0, NAVY)
        r = c0.paragraphs[0].add_run(lab)
        set_run(r, size=11, bold=True, color=WHITE)
        r = c1.paragraphs[0].add_run(val)
        set_run(r, size=12, bold=True, color=INK)
        if val.startswith("["):
            shade(c1, FILL_HINT)
        else:
            shade(c1, VALUE_FILL)
        c0.width = Cm(5.2)
        c1.width = Cm(11.6)

    add_p(doc, "Team members (edit names and roles in this table, or regenerate with --team / --members)", size=10, italic=True, color=MUTED, space_before=10, space_after=4)
    member_rows = [(str(i + 1), n, role) for i, (n, role) in enumerate(members)]
    t = add_table(doc, ["No.", "Full name", "Role"], member_rows, col_widths=[1.6, 8.5, 6.7])
    if placeholder:
        for row in t.rows[1:]:
            for cell in row.cells:
                shade(cell, FILL_HINT)

    add_p(
        doc,
        "How to reprint with your names:  py -3 tools/build_project_report.py --team \"Your team\" --members \"Name One:Lead, Name Two:Curriculum, Name Three:Design\" --event \"Your competition\" --institution \"Your college\"",
        size=9,
        italic=True,
        color=MUTED,
        space_after=12,
    )

    add_p(doc, "Declaration", size=14, bold=True, color=HEADING, space_before=4)
    add_p(
        doc,
        f"We, the members of {team}, declare that this report describes original work on GyanQuest (ImpactX). "
        "Cited ideas from education research are referenced. The software, curriculum design, and this document "
        "are submitted for the competition named above. We confirm that the product can be demonstrated in a browser.",
        align="justify",
    )
    add_p(doc, "Signatures (optional): ______________________     Date: ______________", size=10, color=MUTED)

    # Contents
    add_h(doc, "Contents", 1)
    toc = [
        "1. Abstract",
        "2. Introduction and problem",
        "3. Objectives",
        "4. Users and Bangladesh context",
        "5. Solution overview",
        "6. Pedagogy and learning design",
        "7. Product features",
        "8. Curriculum coverage",
        "9. Technology and architecture",
        "10. Why missions are 2D and 3D is a lab",
        "11. AI tutor (optional, never blocking)",
        "12. Comparison with typical EdTech",
        "13. Methodology",
        "14. Results and current status",
        "15. Impact",
        "16. Limitations and ethics",
        "17. Roadmap",
        "18. Conclusion",
        "19. References",
        "20. Appendix: how to run a demo",
    ]
    for line in toc:
        add_p(doc, line, size=11, space_after=2)

    # 1 Abstract
    add_h(doc, "1. Abstract", 1)
    add_p(
        doc,
        "GyanQuest is a bilingual (English and Bangla), offline-first learning platform that turns school STEM topics "
        "into playable Canvas 2D missions. One shared engine serves 28 subject titles. Learners act on a stage "
        "(drag, tap, predict, myth-bust), get live coach feedback, unlock a digital book, and may ask an optional AI tutor. "
        "Progress is gated by an 80% fluency check before Mastery. A separate 3D Specimen Lab lets learners inspect "
        "scientific models without forcing 3D onto every lesson. Core play works in a browser with no install and no "
        "cloud requirement. Optional Groq chat and a Laravel admin layer add accounts and tutoring when a network exists. "
        "Impact is converting curriculum into mastery-gated missions at the cost of a browser, not a hardware lab.",
        align="justify",
    )
    add_table(
        doc,
        ["Indicator", "Value (August 2026 live catalog)"],
        [
            ["Subject titles in the catalog", "28 (all live folders)"],
            ["Uniqueness-pass games", "16"],
            ["Finished missions", "25"],
            ["Instructional layers", "4 (canvas, coach, book, optional AI)"],
            ["Fluency gate before Mastery", "80%"],
            ["Language modes", "2 (English, Bangla)"],
        ],
        col_widths=[8.5, 8.3],
    )

    # 2 Intro
    add_h(doc, "2. Introduction and problem", 1)
    add_p(
        doc,
        "School STEM in many Bangladeshi classrooms is still lecture-heavy. Tools are fragmented: a video here, "
        "a quiz app there, a 3D demo that only runs on a lab PC. Shared phones, flaky internet, and no dedicated "
        "hardware lab are the real conditions. Learners watch and click multiple-choice items. Teachers cannot "
        "see whether a concept was enacted or only memorized.",
        align="justify",
    )
    add_p(
        doc,
        "GyanQuest attacks that gap. The product is one engine, many subjects, short missions, and a fluency lock. "
        "It is designed so a village computer class and a city coaching center can run the same titles.",
        align="justify",
    )

    # 3 Objectives
    add_h(doc, "3. Objectives", 1)
    numbered(
        doc,
        [
            "Turn stated learning objectives into playable missions with a clear win condition.",
            "Keep the play stage readable: Canvas 2D for graded missions; 3D only where spatial inspection helps.",
            "Restore hub vs in-mission state after reload (save-v2 in the browser).",
            "Offer Bangla and English as first-class modes, not an afterthought overlay.",
            "Keep AI optional: the game must teach if Groq is down or the key is missing.",
            "Stay install-free: a modern browser is the only required hardware.",
        ],
    )

    # 4 Users
    add_h(doc, "4. Users and Bangladesh context", 1)
    add_table(
        doc,
        ["Stakeholder", "What changes"],
        [
            ["Learner (roughly ages 9-16)", "Enacts the idea, fails safely, retries; language toggle EN/BN."],
            ["Teacher / facilitator", "One pattern across 28 titles; coach lines and books, not 28 different apps."],
            ["School / coaching center", "No Unity install, no GPU lab; shared devices still work."],
            ["Competition jury", "A runnable demo: landing, a 2D mission, Specimen Lab, optional tutor."],
        ],
        col_widths=[5.5, 11.3],
    )
    add_p(
        doc,
        "Everyday examples in missions and books are written to feel local (Bangladesh-friendly stories) so transfer "
        "is not locked to a foreign textbook voice.",
        align="justify",
    )

    # 5 Solution
    add_h(doc, "5. Solution overview", 1)
    add_p(
        doc,
        "The learner path is: Landing catalog → subject game → Mission Hub → Canvas 2D mission → coach feedback → "
        "digital book → optional Groq tutor. Progress is stored locally. Login and admin sync exist but are optional.",
        align="justify",
    )
    add_p(doc, "Four instructional layers (canvas leads; everything else follows)", size=12, bold=True, color=HEADING, space_before=6)
    add_table(
        doc,
        ["Layer", "Job", "Pedagogy"],
        [
            ["1. Interactive canvas", "Source of truth. Drag, tap, tune, myth-bust.", "Constructivism; Predict / Observe / Explain"],
            ["2. Game coach", "Live commentary on streaks and misses.", "Formative assessment; Flow; ZPD scaffolding"],
            ["3. Digital book", "Retrieval practice and dual coding; term links.", "Bloom; spaced look-back"],
            ["4. AI chatbot", "On-demand Socratic tutor via a server proxy.", "Cognitive apprenticeship; optional"],
        ],
        col_widths=[4.2, 6.6, 6.0],
    )

    # 6 Pedagogy
    add_h(doc, "6. Pedagogy and learning design", 1)
    add_p(
        doc,
        "Each topic follows a ten-step spiral copied from the Force Fighter pattern: hook, watch, sort, try, explain, "
        "name the rule, stretch, myth-bust, fluency drill, mastery. Difficulty rises inside a topic; it does not jump "
        "to 'level 10 equals the hardest chapter of the whole subject.'",
        align="justify",
    )
    add_p(doc, "Mission cycle (engine pedagogy): Recall → Predict → Act → Feedback → 80% fluency → Mastery.", bold=True, size=11)
    add_p(
        doc,
        "The 80% fluency gate is a product rule, not a slogan. Mastery stays locked until the drill score is earned. "
        "That is the opposite of 'time watched equals progress.' Flow Autopilot today is streak-based scaffolding "
        "(honest limit: not a full adaptive engine).",
        align="justify",
    )

    # 7 Features
    add_h(doc, "7. Product features", 1)
    bullets(
        doc,
        [
            "28 live catalog titles with a shared hub, coach, chips, and save format.",
            "Bilingual UI (EN/BN) including Specimen Lab labels.",
            "Digital mission books with linked terms that can open the tutor.",
            "Unlock-books demo toggle for juries who need to browse books quickly.",
            "Optional landing login (Laravel Sanctum) without blocking offline play.",
            "3D Specimen Lab: orbit, numbered pins, curved callout to the part name and explanation.",
            "Pin editor for placing labels on models; placements can be exported as JSON.",
            "Secure Groq proxy with a chain of free chat models and a local fallback script.",
        ],
    )

    # 8 Curriculum
    add_h(doc, "8. Curriculum coverage", 1)
    add_table(
        doc,
        ["Track", "Titles"],
        [
            ["Core sciences (5)", "Force Fighter, Chemistry Lab, Bio Explorer, Math Quest, Eco Guardian"],
            ["CS and tech (10)", "ICT, Web, Backend, Database/SQL, Networking, Cyber Shield, OS & Hardware, AI Lab, ML Lab, Data Science"],
            ["Engineering (5)", "Electrical, Mechanical, Civil, Electronics & Robotics, Green Tech"],
            ["Advanced science (4)", "Astronomy & Space, Geology & Earth, Human Anatomy, Genetics & Biotech"],
            ["Math extended (4)", "Statistics, Geometry & Trig, Calculus, Discrete Math"],
        ],
        col_widths=[4.5, 12.3],
    )
    add_p(
        doc,
        "Honesty: not every one of the 28×10×10 cells is a unique authored story yet. Force Fighter and several "
        "uniqueness-pass titles carry the deepest labs. Other titles share the same spiral shell and are being filled "
        "subject by subject. The catalog is live; depth is uneven, and this report does not hide that.",
        align="justify",
        italic=True,
    )

    # 9 Tech
    add_h(doc, "9. Technology and architecture", 1)
    add_table(
        doc,
        ["Layer", "Choice", "Why this, not the alternative"],
        [
            ["Gameplay", "HTML5 Canvas 2D + vanilla JS", "No Unity/WebGL download; readable missions on low-end devices"],
            ["Platform", "Custom engine/ (hub, persist, pedagogy)", "One product, not iframes inside an LMS"],
            ["Landing", "HTML/CSS/JS", "Fast catalog, bilingual copy"],
            ["AI tutor", "Python groq_proxy.py + Groq API", "Key never in the browser; games work if chat is down"],
            ["Auth/admin", "Laravel + Sanctum + Filament (optional)", "School-owned admin; SQLite or MySQL"],
            ["Progress", "localStorage save-v2", "Instant restore on flaky school networks"],
            ["3D lab", "Three.js (WebGL) Specimen Lab", "Spatial inspection without forcing 3D onto every lesson"],
        ],
        col_widths=[3.4, 5.4, 8.0],
    )
    add_p(doc, "Architecture (learner in the center)", size=12, bold=True, color=HEADING, space_before=4)
    add_p(
        doc,
        "Learner → Landing + Mission Hub → Shared engine (UI, pedagogy, persistence) → Subject package "
        "(boot, missions, Canvas scenes) → localStorage. Side doors: Groq tutor proxy; optional Sanctum API + Filament.",
        align="justify",
    )

    # 10 2D vs 3D
    add_h(doc, "10. Why missions are 2D and 3D is a lab", 1)
    add_p(
        doc,
        "Graded missions need a stable camera, readable chips, and a coach that never hides behind a mesh. "
        "3D everywhere would raise load time, asset size, and QA cost across 28 subjects. 3D none would waste "
        "the chance to inspect a cell, a heart, or an engine in space.",
        align="justify",
    )
    add_table(
        doc,
        ["", "Canvas 2D missions", "3D Specimen Lab"],
        [
            ["Purpose", "Structured outcomes, fluency, books", "Inspect a model; pin parts; read a callout"],
            ["Success", "Pass the spiral; 80% drill", "Find the organelle / part and name it"],
            ["Device load", "Must stay light on every title", "Acceptable as an optional lab"],
            ["Authoring", "Fast scene + chips + coach", "GLB + pins; not 280 missions of 3D"],
        ],
        col_widths=[3.2, 6.8, 6.8],
    )
    add_p(
        doc,
        "Specimen Lab uses the site’s dark monsoon theme so the 3D stage matches the rest of GyanQuest. "
        "Pins sit on real parts (not only the outer rim). Clicking a pin draws a curved leader to a card with the "
        "part title and a short explanation, generated from that pin’s own bilingual text.",
        align="justify",
    )

    # 11 AI
    add_h(doc, "11. AI tutor (optional, never blocking)", 1)
    add_p(
        doc,
        "Book terms open a chat that posts to /api/chat. The Python proxy keeps GROQ_API_KEY on the server. "
        "Groq retired llama-3.3-70b-versatile for free/developer use (16 August 2026). The proxy now walks a "
        "fallback chain of live free Groq text models: openai/gpt-oss-20b, openai/gpt-oss-120b, qwen/qwen3.6-27b, "
        "groq/compound-mini, groq/compound, openai/gpt-oss-safeguard-20b, allam-2-7b. Rate limits or a dead ID "
        "move to the next model. If Groq is blocked, a local tutor script still answers so class is not stuck.",
        align="justify",
    )
    bullets(
        doc,
        [
            "First tap on a word: full picture (5-7 sentences) and a fixed opt-in question.",
            "Follow-up turns: deeper Socratic question about the idea just taught.",
            "Speech, TTS, and prompt-guard models are not used for this chat path.",
        ],
    )

    # 12 Comparison
    add_h(doc, "12. Comparison with typical EdTech", 1)
    add_table(
        doc,
        ["Dimension", "Typical EdTech", "GyanQuest"],
        [
            ["Engagement", "Watch + MCQ", "Act in a 2D mission"],
            ["Progress", "Time watched / completion %", "80% fluency lock before Mastery"],
            ["Access", "Cloud required", "Offline-first; chat is extra"],
            ["3D", "3D everywhere or none", "3D Specimen Lab only, missions stay 2D"],
            ["Language", "English UI, Bangla as subtitle", "Two strict language modes"],
            ["Subjects", "One app per topic", "One engine, 28 titles"],
        ],
        col_widths=[3.4, 6.7, 6.7],
    )

    # 13 Method
    add_h(doc, "13. Methodology", 1)
    numbered(
        doc,
        [
            "Map NCTB-adjacent school topics to 10-mission shells per subject.",
            "Design the Force Fighter spiral (hook to mastery) as the template for every title.",
            "Build Canvas scenes and persist (save-v2); QA uniqueness of activities where claimed.",
            "Author bilingual copy; keep Bangla in Noto Sans Bengali on the web client.",
            "Isolate 3D in Specimen Lab with pin JSON that teachers can correct.",
            "Keep Groq behind a proxy; test local fallback when the network fails.",
            "Iterate from playtests: coach clarity, pin placement, dark-theme consistency.",
        ],
    )

    # 14 Results
    add_h(doc, "14. Results and current status", 1)
    add_p(
        doc,
        "As of August 2026 the landing catalog lists 28 live titles. Sixteen games pass an internal uniqueness bar. "
        "Twenty-five missions are finished to a playable standard. The remaining mission slots are structured stubs "
        "(same hub, coming-soon cards), which is a capacity choice, not a hidden crash. Specimen Lab ships with a "
        "catalog of classroom models and baked pin positions for animal cell, plant cell, heart, kidney, and city, "
        "plus an editor for the rest.",
        align="justify",
    )
    add_p(
        doc,
        "A jury demo that always works: open the landing page, switch language, play one Force Fighter or Bio Explorer "
        "mission through fluency, open a book term, then open Specimen Lab and tap a numbered pin.",
        align="justify",
    )

    # 15 Impact
    add_h(doc, "15. Impact", 1)
    add_p(
        doc,
        "Impact equals converting school topics into mastery-gated, bilingual, offline missions at the cost of a browser, "
        "not a lab. Learners move from watching to enacting. Teachers get one pattern instead of 28 tools. "
        "Communities with shared phones are not locked out by a 3D-only download.",
        align="justify",
    )
    add_table(
        doc,
        ["Alignment", "How GyanQuest contributes"],
        [
            ["Quality education (SDG 4)", "Active missions, fluency gate, Bangla/English parity"],
            ["Industry and innovation (SDG 9)", "Browser STEM lab; optional self-hosted admin"],
            ["Reduced inequalities (SDG 10)", "Offline-first; low device bar; no paid 3D runtime"],
        ],
        col_widths=[6.0, 10.8],
    )

    # 16 Limits
    add_h(doc, "16. Limitations and ethics", 1)
    bullets(
        doc,
        [
            "Flow Autopilot is streak-based; it is not a full adaptive difficulty model.",
            "Some catalog titles are still shallow relative to Force Fighter.",
            "3D models come from Sketchfab-style classroom assets; credits stay on the lab sheet.",
            "AI answers can be wrong; the canvas and book remain the source of truth.",
            "No biometric or unnecessary personal data is required for core play.",
            "API keys never ship to the browser; chat is off unless the proxy is running.",
        ],
    )

    # 17 Roadmap
    add_h(doc, "17. Roadmap", 1)
    add_table(
        doc,
        ["Horizon", "Work"],
        [
            ["Now", "28 titles, fluency gate, Specimen Lab, Groq fallback chain, optional Laravel"],
            ["Next", "Teacher dashboard on Filament; deeper uniqueness on remaining titles; more baked 3D pins"],
            ["Later", "Limited 3D beats inside selected lessons; richer adaptive scaffolding"],
        ],
        col_widths=[3.5, 13.3],
    )

    # 18 Conclusion
    add_h(doc, "18. Conclusion", 1)
    add_p(
        doc,
        f"{team} submits GyanQuest as a complete, demonstrable learning system: not a slide deck about gamification, "
        "but a browser lab with a pedagogy loop, a fluency lock, bilingual copy, an honest 2D/3D split, and an AI tutor "
        "that cannot hold the class hostage. We ask the jury to play a mission, pin a cell, and switch language. "
        "That is the product.",
        align="justify",
    )

    # 19 References
    add_h(doc, "19. References", 1)
    refs = [
        "Bloom, B. S. (1956). Taxonomy of Educational Objectives. Longman.",
        "Bruner, J. S. (1966). Toward a Theory of Instruction. Harvard University Press.",
        "Csikszentmihalyi, M. (1990). Flow: The Psychology of Optimal Experience. Harper & Row.",
        "Gee, J. P. (2007). What Video Games Have to Teach Us About Learning and Literacy (2nd ed.). Palgrave Macmillan.",
        "Vygotsky, L. S. (1978). Mind in Society. Harvard University Press.",
        "White, R., & Gunstone, R. (1992). Probing Understanding. Falmer Press. (Predict-Observe-Explain).",
        "Groq. (2026). Model deprecations: llama-3.3-70b-versatile shutdown for free/developer tiers (16 August 2026). https://console.groq.com/docs/deprecations",
        "GyanQuest repository and internal design notes: TECHNOLOGY.md, pedagogy engine, Specimen Lab.",
    ]
    for i, ref in enumerate(refs, 1):
        add_p(doc, f"[{i}]  {ref}", size=10, space_after=4)

    # 20 Appendix
    add_h(doc, "20. Appendix: how to run a demo", 1)
    add_p(doc, "Frontend (from the project root)", size=12, bold=True, color=HEADING)
    add_p(doc, "1. Copy .env.example to .env and add GROQ_API_KEY if you want chat.", size=11)
    add_p(doc, "2. py -3 tools/groq_proxy.py", size=11)
    add_p(doc, "3. Open http://127.0.0.1:5500/", size=11)
    add_p(doc, "Optional API: cd backend → composer install → php artisan serve (http://127.0.0.1:8000/).", size=11)
    add_p(
        doc,
        "Without chat: py -3 -m http.server 5500 still runs every mission and Specimen Lab. "
        "Progress lives in the browser until login sync is used.",
        align="justify",
    )
    add_p(
        doc,
        "End of report. Cream cells on the cover (if present) are fill-ins: replace them in Word or regenerate the file.",
        italic=True,
        size=10,
        color=MUTED,
        space_before=12,
    )

    out = Path(args.out) if args.out else OUT
    if not out.is_absolute():
        out = ROOT / out
    try:
        doc.save(out)
    except PermissionError:
        alt = out.with_name(out.stem + "_updated" + out.suffix)
        doc.save(alt)
        print(f"Original file is open. Wrote {alt} instead.")
        return alt
    return out


def main():
    ap = argparse.ArgumentParser(description="Write GyanQuest_Project_Report.docx")
    ap.add_argument("--team", default="Team Alpha", help="Team name printed on the cover")
    ap.add_argument(
        "--members",
        default="Md Tahsan Islam:Lead and platform, Mufrid Johanee:Curriculum and missions, Shadiya Zaman Tanha:UX and content, Md Mohaimenul Islam:Systems and pedagogy",
        help="Comma-separated names. Optional Role after a colon. Example: 'Ayesha:Lead, Rahim:Design'",
    )
    ap.add_argument("--event", default="", help="Competition or event name (yellow fill-in if omitted)")
    ap.add_argument("--institution", default="", help="College / school / organization")
    ap.add_argument("--out", default="", help="Output .docx path")
    args = ap.parse_args()
    path = build(args)
    print(f"Wrote {path}")


if __name__ == "__main__":
    main()
